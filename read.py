from docx import Document
from docx.table import Table
import re
import argparse
import os
from tqdm import tqdm

global table_cnt

def read_table(table):
    data = []
    keys = None
    for _rowid, row in enumerate(table.rows):
        text = [' '+cell.text+' ' for cell in row.cells]
        # # delete duplicate merged cells
        # for i, t in enumerate(text):
        #     if i == 0:
        #         continue
        #     if t == text[i-1]:
        #         text[i] = ' '
        text = [re.sub('[\n+\t]', ' ', t) for t in text]
        data.append('|'+'|'.join(text)+'|')
        if _rowid == 0:
            data.append('|'+'|'.join([' '+'-'*4+' ' for _ in range(len(text))])+'|')
    return '###'.join(data)

def whitespace(text):
    if text in ['\n', '\t', ' ', '']:
        return True
    return False

def istable(text):
    prog = re.compile(r'^表\d+ ')
    if re.match(prog, text):
        return True
    return False

def read_paragraph(paragraph, doc):
    data = []
    table_cnt = 0
    for para in paragraph:
        text = para.text
        if not whitespace(text):
            data.append(text)
            
            # append table
            if istable(text):
                text = read_table(doc.tables[table_cnt])
                table_cnt += 1
                data.append(text)
            
    return '\n'.join(data)

def mkdirs(path):
    if not os.path.exists(path):
        os.makedirs(path, exist_ok=True)

def write_file(path, text):
    with open(path, 'w', encoding='utf-8') as f:
        f.write(text)

def parse_args():
    parser = argparse.ArgumentParser(description='read docx file')
    parser.add_argument('--docx', type=str, default='test.docx', help='docx file path')
    parser.add_argument('--docdir', type=str, default='./', help='docx file dir')
    parser.add_argument('--output', type=str, default='../txt', help='output file dir')
    args = parser.parse_args()
    return args

if __name__ == '__main__':
    args = parse_args()
    
    print('parse doc files:')
    # read docx file
    if args.docdir is not None:
        docdir = args.docdir
        for docname in tqdm(os.listdir(docdir)):
            if not docname.endswith(('.docx', '.doc')):
                continue
            docpath = os.path.join(docdir, docname)
            doc = Document(docpath)
            text = read_paragraph(doc.paragraphs, doc)
            output_dir = args.output
            mkdirs(output_dir)
            write_file(os.path.join(output_dir, docname.split('.')[0]+'.txt'), text)
    else:
        doc = Document('../附件3——信用风险权重法表内资产风险权重表外项目信用转换系数及合格信用风险缓释工具.docx')
        # doc = Document(args.docx)

        # import pdb;pdb.set_trace()

        # text = read_table(doc.tables[0])
        # print(text)
        
        ###
        text = read_paragraph(doc.paragraphs, doc)
        print(text)
        
        output_dir = args.output
        mkdirs(output_dir)
        
        write_file(os.path.join(output_dir, 'test.txt'), text)
        
        # text = read_table(doc.tables[-1])
        # print(text)

        # import pdb;pdb.set_trace()
    print('finish')